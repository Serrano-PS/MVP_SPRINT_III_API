from flask_openapi3 import OpenAPI, Info, Tag
from flask import redirect
from urllib.parse import unquote
from flask import Flask, request, send_from_directory, render_template, make_response, send_file, Response
from sqlalchemy.exc import IntegrityError
import io
# Exportação CSV
import csv
# Exportação PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
# Exportação Xlsx
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
# Exportação Docx
from docx import Document
import requests


from model import Session, Valve
from logger import logger
from schemas import *
from flask_cors import CORS

info = Info(title="MVP API", version="1.0.0")
app = OpenAPI(__name__, info=info)
CORS(app)

# definindo tags para categorizar as operações da API

# a tag 'home_tag' é para documentação em geral, especificamente para a seleção de uma ferramenta de documentação chamada Swagger.
home_tag = Tag(name="Documentação",
               description="Seleção de documentação: Swagger")

# a tag 'valve_tag' é para ações relacionadas a válvulas, como adicionar, visualizar, editar e remover válvulas da base de dados.
valve_tag = Tag(
    name="Válvula", description="Adição, visualização,edição e remoção de válvulas da base")


def home():
    """Redireciona para a tela que permite a escolha do estilo de documentação Swagger.
    """
    return redirect('/openapi/swagger')


"""Adiciona a rota '/' para a função home com método GET
    """
app.add_url_rule('/', view_func=home, methods=['GET'])


@app.route('/favicon.ico')
def favicon():
    return send_from_directory('static', 'favicon.ico', mimetype='image/x-icon')


@app.post('/valve', tags=[valve_tag],
          responses={"200": ValveViewSchema, "409": ErrorSchema, "400": ErrorSchema})
def add_valve(form: ValveSchema):
    """Adiciona uma nova Válvula à base de dados.

    Retorna uma representação das válvulas.
    """
    # Criando objeto da classe Valve
    valve = Valve(
        nome=form.nome,
        descricao=form.descricao,
        tipo=form.tipo,
        vazao=form.vazao)
    logger.debug(f"Adicionando válvula de nome: '{valve.nome}'")
    try:
        # criando conexão com a base
        session = Session()
        # adicionando válvula
        session.add(valve)
        # efetivando o comando de adição de novo item na tabela
        session.commit()
        logger.debug(f"Adicionada válvula de nome: '{valve.nome}'")
        return apresenta_valve(valve), 200

    except IntegrityError as e:
        # como a duplicidade do nome é a provável razão do IntegrityError
        error_msg = "Válvula de mesmo nome já salva na base :/"
        logger.warning(
            f"Erro ao adicionar válvula '{valve.nome}', {error_msg}")
        return {"message": error_msg}, 409

    except Exception as e:
        # caso um erro fora do previsto
        error_msg = "Não foi possível salvar novo item :/"
        logger.warning(
            f"Erro ao adicionar válvula '{valve.nome}', {error_msg}")
        return {"message": error_msg}, 400


@app.get('/valves', tags=[valve_tag],
         responses={"200": ListagemValveSchema, "404": ErrorSchema})
def get_valves():
    """Faz a busca por todas as Válvulas cadastradas

    Retorna uma representação da listagem de válvulas.
    """
    logger.debug(f"Coletando Válvulas ")
    # criando conexão com a base
    session = Session()
    # fazendo a busca
    valves = session.query(Valve).all()

    if not valves:
        # se não há produtos cadastrados
        return {"válvulas": []}, 200
    else:
        logger.debug(f"%d válvulas econtradss" % len(valves))
        # retorna a representação de produto
        print(valves)
        return apresenta_valves(valves), 200


@app.delete('/valve', tags=[valve_tag],
            responses={"200": ValveDelSchema, "404": ErrorSchema})
def del_produto(query: ValveBuscaSchema):
    """Remove uma Válvula da base de dados.

      Retorna 204 em caso de sucesso.
      """
    valve_nome = unquote(unquote(query.nome))
    print(valve_nome)
    logger.debug(f"Deletando dados sobre válvula #{valve_nome}")
    # criando conexão com a base
    session = Session()
    # fazendo a remoção
    count = session.query(Valve).filter(
        Valve.nome == valve_nome).delete()
    session.commit()

    if count:
        # retorna a representação da mensagem de confirmação
        logger.debug(f"Deletada válvula #{valve_nome}")
        return {"mesage": "Válvula removida", "Nome": valve_nome}
    else:
        # se o produto não foi encontrado
        error_msg = "Válvula não encontrada na base :/"
        logger.warning(
            f"Erro ao deletar válvula #'{valve_nome}', {error_msg}")
        return {"mesage": error_msg}, 404


@app.put('/valve', tags=[valve_tag], responses={"200": ValveViewSchema, "404": ErrorSchema})
def update_valve(form: ValveUpdateSchema):
    """Atualiza as informações de uma Válvula a partir do ID da válvula.

    Retorna uma representação das válvulas.
    """
    valve_id = form.id
    logger.debug(f"Atualizando dados sobre válvula #{valve_id}")
    # criando conexão com a base
    session = Session()
    # fazendo a busca
    valve = session.query(Valve).filter(Valve.id == valve_id).first()

    if not valve:
        # se a válvula não foi encontrada
        error_msg = "Válvula não encontrada"
        logger.warning(f"{error_msg}: {valve_id}")
        return {"message": error_msg}, 404

    # atualizando os dados da válvula
    valve.nome = form.nome
    valve.descricao = form.descricao
    valve.tipo = form.tipo
    valve.vazao = form.vazao

    # efetivando a atualização
    session.commit()

    logger.debug(f"Dados da válvula '{valve.nome}' atualizados")
    return apresenta_valve(valve), 200


@app.get('/export/csv', tags=[valve_tag])
def export_to_csv():
    """Exportar os dados para um aquivo csv.
    """
    session = Session()
    valves = session.query(Valve).all()
    output = io.StringIO()  # Cria um buffer de string para o CSV
    csv_writer = csv.writer(output)

    # csv_writer = csv.writer(output)
    # Escreve o cabeçalho do CSV
    csv_writer.writerow(['ID', 'Nome', 'Descrição', 'Tipo', 'Vazão'])
    # Escreve os dados das válvulas no CSV
    for valve in valves:
        csv_writer.writerow(
            [valve.id, valve.nome, valve.descricao, valve.tipo, valve.vazao])

    # Obtém os dados do CSV como string
    csv_data = output.getvalue()

    # Cria a resposta com o CSV
    response = make_response(csv_data)
    response.headers["Content-Disposition"] = "attachment; filename=valves.csv"
    response.headers["Content-type"] = "text/csv; charset=utf-8"
    return response

# Definir rota para exportar para PDF


@app.get('/export/pdf', tags=[valve_tag])
def export_to_pdf():
    """Exportar os dados para um aquivo pdf.
    """
    session = Session()
    valves = session.query(Valve).all()

    # Criar um buffer de bytes para o PDF
    buffer = io.BytesIO()

    # Criar um documento PDF usando ReportLab
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    elements = []

    # Criar tabela para os dados
    data = [['ID', 'Nome', 'Descrição', 'Tipo', 'Vazão']]
    for valve in valves:
        data.append([valve.id, valve.nome, valve.descricao,
                    valve.tipo, valve.vazao])

    table = Table(data)
    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])
    table.setStyle(style)
    elements.append(table)

    # Adicionar a tabela ao documento PDF
    doc.build(elements)

    # Retornar o PDF como resposta
    response = make_response(buffer.getvalue())
    response.headers["Content-Disposition"] = "attachment; filename=valves.pdf"
    response.headers["Content-type"] = "application/pdf"

    return response


@app.get('/export/xml', tags=[valve_tag])
def export_to_xml():
    """Exportar os dados para um aquivo xml.
    """
    try:
        session = Session()
        valves = session.query(Valve).all()

        xml_data = '<?xml version="1.0" encoding="UTF-8"?>\n<valves>\n'
        for valve in valves:
            xml_data += f'    <valve>\n        <id>{valve.id}</id>\n        <nome>{valve.nome}</nome>\n        <descricao>{valve.descricao}</descricao>\n        <tipo>{valve.tipo}</tipo>\n        <vazao>{valve.vazao}</vazao>\n    </valve>\n'
        xml_data += '</valves>'

        response = make_response(xml_data)
        response.headers["Content-Disposition"] = "attachment; filename=valves.xml"
        response.headers["Content-type"] = "application/xml; charset=utf-8"
        return response
    except Exception as e:
        return {"message": "Erro ao exportar para XML"}, 500


@app.get('/export/xlsx', tags=[valve_tag])
def export_to_xlsx():
    """Exportar os dados para um aquivo xlsx.
    """
    session = Session()
    valves = session.query(Valve).all()

    # Criando um arquivo XLSX em memória
    output = io.BytesIO()
    workbook = openpyxl.Workbook()
    worksheet = workbook.active

    # Aplicando estilos de formatação à primeira linha (cabeçalho)
    header_font = Font(bold=True)
    header_alignment = Alignment(horizontal='center')
    header_fill = PatternFill(start_color='DDDDDD',
                              end_color='DDDDDD', fill_type='solid')
    header_border = Border(bottom=Side(style='thin'))
    for col_num, header in enumerate(['ID', 'Nome', 'Descrição', 'Tipo', 'Vazão'], 1):
        cell = worksheet.cell(row=1, column=col_num, value=header)
        cell.font = header_font
        cell.alignment = header_alignment
        cell.fill = header_fill
        cell.border = header_border

    # Preenchendo os dados das válvulas na planilha
    for row_num, valve in enumerate(valves, 2):
        for col_num, value in enumerate([valve.id, valve.nome, valve.descricao, valve.tipo, valve.vazao], 1):
            cell = worksheet.cell(row=row_num, column=col_num, value=value)
            cell.border = Border(top=Side(style='thin'),
                                 bottom=Side(style='thin'))

    # Definindo largura das colunas
    for col in worksheet.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        worksheet.column_dimensions[column].width = adjusted_width

    # Salvando o arquivo XLSX em memória
    workbook.save(output)
    output.seek(0)

    # Criando a resposta com o arquivo XLSX
    response = make_response(output.read())
    response.headers["Content-Disposition"] = "attachment; filename=valves.xlsx"
    response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    return response


@app.get('/export/docx', tags=[valve_tag])
def export_to_docx():
    """Exportar os dados para um aquivo docx.
    """
    session = Session()
    valves = session.query(Valve).all()

    # Criando um documento DOCX em memória
    document = Document()

    # Adicionando título
    document.add_heading('Exportação de Dados de Válvulas', level=1)

    # Adicionando tabela
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = True

    # Preenchendo o cabeçalho da tabela
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ID'
    header_cells[1].text = 'Nome'
    header_cells[2].text = 'Descrição'
    header_cells[3].text = 'Tipo'
    header_cells[4].text = 'Vazão'

    # Preenchendo os dados das válvulas na tabela
    for valve in valves:
        row_cells = table.add_row().cells
        row_cells[0].text = str(valve.id)
        row_cells[1].text = valve.nome
        row_cells[2].text = valve.descricao
        row_cells[3].text = valve.tipo
        row_cells[4].text = str(valve.vazao)

    # Salvando o documento DOCX em memória
    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    # Criando a resposta com o documento DOCX
    response = make_response(output.read())
    response.headers["Content-Disposition"] = "attachment; filename=valves.docx"
    response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response


RAPIDAPI_KEY = "9c1568c93emshab13bdece0f330dp1b9174jsn21d163e3c690"


@app.post('/export/URL_to_PDF', tags=[valve_tag])
def export_URL_to_PDF():
    """Exportar uma pagina web (URL) para um aquivo pdf.
    """
    url_to_pdf_api = "https://html2pdf-rocket.p.rapidapi.com/pdf"

    headers = {
        "content-type": "application/json",
        "X-RapidAPI-Key": RAPIDAPI_KEY,
        "X-RapidAPI-Host": "html2pdf-rocket.p.rapidapi.com"
    }

    payload = {"value": "http://simdut.com.br"}

    response = requests.post(url_to_pdf_api, json=payload, headers=headers)

    if response.status_code == 200:
        pdf_content = response.content
        response = Response(pdf_content, content_type='application/pdf')
        response.headers['Content-Disposition'] = 'attachment; filename=exported_URL.pdf'
        return response
    else:
        return {"error": "Failed to generate PDF"}, 500


# TENTEI OBTER URL PASSADA PELO FRONT-END, MAS ELE GERA O PDF APENAS COM O LINK. NÃO ENCONTREI O MOTIVO DO ERRO.
# def export_URL_to_PDF():
#     data = request.json  # Obtém os dados JSON do corpo da solicitação

#     if 'value' in data:
#         url = data['value']  # Obtém a URL passada pelo front-end

#         url_to_pdf_api = "https://html2pdf-rocket.p.rapidapi.com/pdf"

#         headers = {
#             "content-type": "application/json",
#             "X-RapidAPI-Key": RAPIDAPI_KEY,
#             "X-RapidAPI-Host": "html2pdf-rocket.p.rapidapi.com"
#         }

#         payload = {"value": url}

#         response = requests.post(url_to_pdf_api, json=payload, headers=headers)

#         if response.status_code == 200:
#             pdf_content = response.content
#             response = Response(pdf_content, content_type='application/pdf')
#             response.headers['Content-Disposition'] = 'attachment; filename=exported.pdf'
#             return response
#         else:
#             return {"error": "Failed to generate PDF"}, 500
#     else:
#         return {"error": "A URL must be provided in the 'value' field"}, 400
